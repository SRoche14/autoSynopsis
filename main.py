from tkinter import *
from tkinter import filedialog
import os
import pandas as pd
from docx import Document
import numpy as np
import re
from datetime import datetime
import math

root = Tk()
root.iconbitmap('kermit_icon.ico')
app_width = 1000
app_height = 1200
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()

x = (screen_width / 2) - (app_width / 2)
y = (screen_height / 2) - (app_height / 2)
root.geometry(f'{app_width}x{app_height}+{int(x)}+{int(y)}')


endLabel = Label(root, text="Completed Making Synopses. Have fun with your new free time :)", font=("Arial", 16))
my_label = Label(root, text="Lists created, now generate synopses!", font=("Arial", 20))

folder2 = "NULL"

column_names = []
table_names = []
row_words = []
logic_list = []
statement_list = []
num2 = 0


def open_conf(label):
    global column_names
    global table_names
    global row_words
    global logic_list
    global statement_list
    root.filename = filedialog.askopenfilename(initialdir="/Users/sroche/Documents/AutoSynopsis",
                                               title="Select Configuration File",
                                               filetypes=(("docx files", "*.docx"), ("all files", "*.*")))

    # works for only 1 header tables
    document = Document(root.filename)
    my_label.pack_forget()
    endLabel.pack_forget()
    table_count = 0
    for table in document.tables:
        data = [[cell.text for cell in row.cells] for row in table.rows]
        df = pd.DataFrame(data)
        df = df.rename(columns=df.iloc[0]).drop(df.index[0]).reset_index(drop=True)
        table_count += 1
        try:
            if table_count == 1:
                find_col_name = [col for col in df.columns if 'Column Name' in col]
                row_word = [col for col in df.columns if 'Row contain' in col]
                table_names = df['Table Name'].tolist()
                table_names = [each_string.lower() for each_string in table_names]
                column_names = df[find_col_name].values.tolist()
                store = []
                for item in column_names:
                    temp = []
                    for name in item:
                        temp.append(name.lower())
                    store.append(temp)
                column_names = store
                row_words = df[row_word].values.tolist()
                store = []
                for item in row_words:
                    temp = []
                    for name in item:
                        temp.append(name.lower())
                    store.append(temp)
                row_words = store
            elif table_count >= 2:
                logic_list_add = df['Logic'].tolist()
                logic_list.append(logic_list_add)
                statement_list_add = df['Statement'].tolist()
                statement_list.append(statement_list_add)

        except ReferenceError:
            print("not here")
        table_count += 1

    label.pack()


def list_maker(array):
    if "=" in array:
        symbol = "="
    elif ">" in array:
        symbol = ">"
    elif "<" in array:
        symbol = "<"
    elif ">=" in array:
        symbol = ">="
    elif "<=" in array:
        symbol = "<="
    else:
        symbol = 'NULL'
    try:
        index = array.index(symbol)
        group1 = array[0:index]
        group2 = array[index + 1:]
        return symbol, group1, group2
    except KeyError:
        return "Null", "Null", "Null"


def string_func(ele):
    return str(ele)


def make_word_doc(output):
    global num2
    global folder2
    num2 += 1

    if folder2 == "NULL":
        root.folder2 = filedialog.askdirectory(initialdir="/Users/sroche/Documents/AutoSynopsis",
                                          title="Select Ouput Folder")
        folder2 = root.folder2
    document = Document()
    document.add_heading('Outputed Synopses', 0)
    document.add_heading('Statements:', level=1)
    for statement in output:
        p = document.add_paragraph(statement)
    name_docx = folder2 + "/completedSynopses" + str(num2) + ".docx"
    document.save(name_docx)
    endLabel.pack()


def develop_sentences(output1_arr, output2_arr):
    output_sentences = []
    log_set = -1
    for logic_sets in logic_list:

        small_set = -1
        log_set += 1
        for small_log_set in logic_sets:
            small_set += 1

            logic_arr = small_log_set.split(',')
            log_holder = []
            current_fulfilled = 0
            for log_statement in logic_arr:
                if log_statement == '':
                    continue
                elif log_statement.find("If") != -1 or log_statement.find("if") != -1:
                    index = int(log_statement.find('{'))
                    end = int(log_statement.find('}') + 1)
                    section = log_statement[index:end]
                    section = section.replace("{", "").replace("}", "")

                    if "^" in section:
                        section = int(section.replace("^", "")) - 1

                        if section in output2_arr.keys():
                            current_fulfilled += 1
                        else:
                            break
                    else:
                        try:
                            section = int(section) - 1
                            if section in output1_arr.keys():
                                current_fulfilled += 1
                        except KeyError:
                            break
                elif ">" or "<" or "=" or "<=" or ">=" in log_statement:
                    life_expectancy = False

                    s = re.split('{|}', log_statement)
                    count = -1

                    for item in s:
                        count += 1
                        if count % 2 == 0:
                            continue
                        elif count % 2 == 1:
                            if "^" in item:
                                output_use = output2_arr
                                item = item.replace("^", "")
                            else:
                                output_use = output1_arr
                            if "#" in item:
                                life_expectancy = True
                                item = item.replace("#", "")
                            try:
                                piece = int(item) - 1

                                val = output_use[piece].replace("$", '').replace(',', '')
                                val = val.replace('/mo', '').replace('/yr', '').replace('/wk', '')
                                if life_expectancy:
                                    year = val.split(' ')
                                    year = int(year[-1])
                                    current_year = datetime.now().year
                                    val = year - current_year
                                s[count] = val
                            except KeyError:
                                print('key error')
                                break
                    s = list(map(string_func, s))

                    try:
                        output = ''.join(s)
                        output = output.replace("]", "").replace("[", "")

                        result = round(eval(output), 2)

                        if result:
                            current_fulfilled += 1
                    except SyntaxError:
                        s = ''.join(s).strip()
                        s = s.split(" ")

                        if s[0] == s[2]:
                            current_fulfilled += 1
                        else:
                            print('not equal')

                if current_fulfilled == len(logic_arr):

                    # count = 0
                    statement = statement_list[log_set][small_set]

                    correct = 0
                    s = re.split('\[|]', statement)
                    count = -1
                    for ele in s:
                        life_expectancy = False
                        dollar_format = False
                        count += 1
                        if count % 2 == 0:
                            continue
                        elif count % 2 == 1:
                            print(ele)

                            tracker = 0
                            if "$" in ele:
                                dollar_format = True
                                ele = ele.replace("$", "")

                            s2 = re.split('{|}', ele)
                            s2 = ' '.join(s2).split()
                            count2 = -1
                            print(s2)
                            error_count = 0
                            for item in s2:
                                count2 += 1
                                try:
                                    if "^" in item:
                                        output_use = output2_arr
                                        item = item.replace("^", "")
                                    else:
                                        output_use = output1_arr
                                    if "#" in item:
                                        life_expectancy = True
                                        item = item.replace("#", "")
                                    try:
                                        piece = int(item) - 1

                                        val = output_use[piece].replace("$", '').replace(',', '')
                                        val = val.replace('/mo', '').replace('/yr', '').replace('/wk', '')
                                        if life_expectancy:
                                            year = val.split(' ')
                                            year = int(year[-1])
                                            current_year = datetime.now().year
                                            val = year - current_year
                                        s2[count2] = val
                                        tracker += 1
                                    except KeyError:
                                        print('key error')
                                        break
                                except ValueError:
                                    error_count += 1
                                    continue
                            s2 = list(map(string_func, s2))
                            output = ''.join(s2)
                            output = output.replace("]", "").replace("[", "")

                            if not math.floor(len(s2) / 2) > error_count and error_count > 0:
                                output = round(eval(output), 2)
                                if dollar_format:
                                    print('here')
                                    dollar = "${:,}".format(int(output))
                                    if '-' in dollar:
                                        dollar = dollar.replace('-', '')
                                        dollar = "-" + dollar
                                    output = dollar
                                s[count] = output
                                correct += 1
                            elif error_count == 0:
                                if dollar_format:
                                    print('here')
                                    dollar = "${:,}".format(int(output))
                                    if '-' in dollar:
                                        dollar = dollar.replace('-', '')
                                        dollar = "-" + dollar
                                    output = dollar
                                s[count] = output
                    s = list(map(string_func, s))

                    try:
                        output = ''.join(s)
                        output = output.replace("]", "").replace("[", "")
                        print(output)
                        if '{' and '}' not in output:
                            output_sentences.append(output)
                    except:
                        print('failure')
                    # while '{' in statement and count <= 20:
                    #     count += 1
                    #     index = int(statement.find('{'))
                    #     end = int(statement.find('}') + 1)
                    #     section = statement[index:end]
                    #     key = section.replace('{', "").replace('}', "")
                    #     operation_arr = ["+", "-", "*", "/"]
                    #     ops = {"+": operator.add, "-": operator.sub, "*": operator.mul, "/": operator.truediv}
                    #     if "$" in key:
                    #         dollar_format = True
                    #         key = key.replace("$", "")
                    #     else:
                    #         dollar_format = False
                    #     if "#" in key:
                    #         life_expectancy = True
                    #         key = key.replace("#", "")
                    #     else:
                    #         life_expectancy = False
                    #     pieces = re.split(r"([*+\-/])", key)
                    #     count2 = -1
                    #     result = 0
                    #     loop_counter = 0
                    #     for data in pieces:
                    #         count2 += 1
                    #         if data in operation_arr and len(pieces) != 1:
                    #             first = pieces[count2 - 1]
                    #             second = pieces[count2 + 1]
                    #             if '^' in first:
                    #                 first = output2_arr[int(first.replace('^', "")) - 1]
                    #             else:
                    #                 first = output1_arr[int(first) - 1]
                    #             first = first.replace('$', "").replace(',', '')
                    #             if '^' in second:
                    #                 second = output2_arr[int(second.replace('^', "")) - 1]
                    #             else:
                    #                 second = output1_arr[int(second) - 1]
                    #             second = second.replace('$', "").replace(',', '')
                    #             if loop_counter == 0:
                    #                 loop_counter += 1
                    #                 if data == '/':
                    #                     result += int(first) / int(second)
                    #                 else:
                    #                     result += int(ops[data](int(first), int(second)))
                    #             else:
                    #                 loop_counter += 1
                    #                 if data == '/':
                    #                     result += int(result) / int(second)
                    #                 else:
                    #                     result += int(ops[data](int(result), int(second)))
                    #         elif len(pieces) == 1:
                    #             if '^' in data:
                    #                 result = output2_arr[int(data.replace('^', "")) - 1]
                    #             else:
                    #                 result = output1_arr[int(data) - 1]
                    #     if dollar_format:
                    #         dollar = "${:,}".format(result)
                    #         if '-' in dollar:
                    #             dollar = dollar.replace('-', '')
                    #             dollar = "-" + dollar
                    #         statement = statement.replace(section, str(dollar))
                    #     elif life_expectancy:
                    #         current_year = datetime.now().year
                    #         result = int(result) - int(current_year)
                    #         statement = statement.replace(section, str(result))
                    #     else:
                    #         statement = statement.replace(section, str(result).title())
                    #     if '{' and '}' not in statement:
                    #         output_sentences.append(statement)
                    # if '{' not in statement and count == 0:
                    #     output_sentences.append(statement)

    make_word_doc(output_sentences)


def gen_synopses():

    if isinstance(column_names, list) and isinstance(table_names, list) and isinstance(row_words, list):
        if len(logic_list) >= 0 and len(statement_list) >= 0:
            root.folder = filedialog.askdirectory(initialdir="/Users/sroche/Documents/AutoSynopsis",
                                                  title="Select Synopses Folder")
            docx_list = os.listdir(root.folder)
            for doc in docx_list:
                output1_list = {}
                output2_list = {}
                table_name_use = []

                for i in table_names:
                    table_name_use.append(i)
                column_name_use = []
                for i in column_names:
                    column_name_use.append(i)
                table_names_docx = []
                if doc == '.DS_Store' or '~' in doc:
                    continue
                file = root.folder + "/" + doc
                document = Document(file)
                for para in document.paragraphs:
                    if para.style.name == 'Detail - Heading Synopsis':
                        table_names_docx.append(para.text.strip().lower())

                count = -1
                information_arr = []
                index_arr = []
                remove_arr = []

                for table_title in table_names:
                    count += 1
                    if table_title == '':
                        continue

                    table_num = table_names_docx.index(table_title)
                    table = document.tables[table_num]
                    data = [[cell.text for cell in row.cells] for row in table.rows]
                    store = []
                    for item in data:
                        temp = []
                        for name in item:
                            temp.append(name.lower())
                        store.append(temp)
                    data = store
                    df = pd.DataFrame(data)
                    df = df.rename(columns=df.iloc[0]).drop(df.index[0]).reset_index(drop=True)
                    try:
                        row_word = row_words[count][0]
                        addition = []
                        mask = np.column_stack([df[col].str.contains(row_word, na=False) for col in df])
                        find_result = np.where(mask == True)
                        if find_result:
                            try:
                                result = [find_result[0][0], find_result[1][0]]
                                addition.append(result)
                                index_arr.append(count)
                            except:
                                remove_arr.append(count)
                                # remove_arr.append(table_title)
                                continue
                            try:
                                result2 = [find_result[0][1], find_result[1][1]]
                                addition.append(result2)
                                information_arr.append(addition)
                            except:
                                information_arr.append(addition)

                    except:
                        continue
                iteration = -1

                for thing in remove_arr:
                    column_name_use[thing] = ''
                    table_name_use[thing] = ''

                column_name_use.remove('')
                table_name_use.remove('')
                for table_title in table_name_use:
                    iteration += 1
                    if table_title == '':
                        continue
                    table_num = table_names_docx.index(table_title)
                    table = document.tables[table_num]
                    data = [[cell.text for cell in row.cells] for row in table.rows]
                    store = []
                    for item in data:
                        temp = []
                        for name in item:
                            temp.append(name.lower())
                        store.append(temp)
                    data = store
                    df = pd.DataFrame(data)
                    df = df.rename(columns=df.iloc[0]).drop(df.index[0]).reset_index(drop=True)
                    group = information_arr[iteration]
                    index_item = index_arr[iteration]
                    for piece in group:
                        row = piece[0]
                        get_cols = column_name_use[iteration][0].split(',')
                        column_label1 = get_cols[0].strip()
                        if len(get_cols) > 1:
                            column_label2 = get_cols[1].strip()
                        else:
                            column_label2 = 'Not existent'
                        # if column_label1 in df:
                        col_list = list(df.columns.values)
                        if any(column_label1 in string for string in col_list):
                            strings = [string for string in col_list if column_label1 in string]
                            column_label1 = strings[0]
                            try:
                                output1 = df.loc[df.index[row], column_label1]
                                adding = output1.strip()
                                output1_list[index_item] = adding
                            except:
                                continue
                            if any(column_label2 in string for string in col_list):
                                strings = [string for string in col_list if column_label2 in string]
                                column_label2 = strings[0]
                                try:
                                    output2 = df.loc[df.index[row], column_label2]
                                    push = output2.strip()
                                    output2_list[index_item] = push
                                except:
                                    continue

                develop_sentences(output1_list, output2_list)


title = Label(root, text="Welcome to AutoSynopsis!", font=("Arial", 25))
title.pack()


conf_btn = Button(root, text="Get configuration file",
                  command=lambda: open_conf(my_label),
                  relief=SUNKEN, padx=20, pady=10,
                  font=("Arial", 16))
conf_btn.pack()

gen_syn_btn = Button(root, text="Generate Synopsis", command=lambda: gen_synopses(),
                     relief=SUNKEN, padx=20, pady=10,
                     font=("Arial", 16))
gen_syn_btn.pack()

start = Label(root, text="Please configure this before generating Synopses.", font=("Arial", 20))
start.pack()

root.mainloop()
